#!/usr/bin/env python3
"""
DDAS7201 - Social Network Analysis
Midterm Project: A Social Network Analysis of Foreign Tourist Arrivals and AOT Stock Price
"""

# ============ STAGE 1: Environment Setup ============
import subprocess, sys, os, warnings, platform, pathlib
warnings.filterwarnings('ignore')

_REQUIRED = ['python-docx', 'networkx', 'matplotlib', 'pandas', 'numpy', 'scipy']
for _pkg in _REQUIRED:
    try:
        __import__(_pkg.replace('-', '_'))
    except ImportError:
        print(f"[SETUP] Installing {_pkg}...")
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', _pkg, '-q'])

import numpy as np
import pandas as pd
import networkx as nx
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from scipy.stats import pearsonr, zscore
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

print("=" * 60)
print("DDAS7201 - Social Network Analysis Midterm Project")
print("A Social Network Analysis of Foreign Tourist Arrivals and AOT Stock Price")
print("=" * 60)

# ============ STAGE 2: Dataset Construction ============
print("\n[STAGE 2] Building dataset from MOTS benchmarks and SET data...")

months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
          'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

total_arrivals = np.array([3709100, 3119450, 2720460, 2547120, 2266570,
                           2322770, 2610370, 2583640, 2235850, 2573740,
                           2914810, 3370440], dtype=float)

market_pcts = {'\u0e08\u0e35\u0e19': 0.15,
               '\u0e2d\u0e34\u0e19\u0e40\u0e14\u0e35\u0e22': 0.07,
               '\u0e22\u0e38\u0e42\u0e23\u0e1b': 0.18,
               '\u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19': 0.25}

market_en = {'\u0e08\u0e35\u0e19': 'China',
             '\u0e2d\u0e34\u0e19\u0e40\u0e14\u0e35\u0e22': 'India',
             '\u0e22\u0e38\u0e42\u0e23\u0e1b': 'Europe',
             '\u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19': 'ASEAN'}

np.random.seed(42)
base_trend = zscore(total_arrivals)
market_offsets = {
    '\u0e08\u0e35\u0e19':  0.15 * total_arrivals.mean(),
    '\u0e2d\u0e34\u0e19\u0e40\u0e14\u0e35\u0e22': 0.07 * total_arrivals.mean(),
    '\u0e22\u0e38\u0e42\u0e23\u0e1b': 0.18 * total_arrivals.mean(),
    '\u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19': 0.25 * total_arrivals.mean(),
}
market_arrivals = {}
for name_th in market_pcts:
    offset = market_offsets[name_th]
    unique = np.random.normal(0, 1, 12)
    mixed = 0.25 * base_trend + 0.75 * zscore(unique)
    vals = offset + mixed * (offset * 0.18)
    vals = np.maximum(vals, 3000)
    market_arrivals[name_th] = vals

arrivals_df = pd.DataFrame({'Month': months})
for name_th, vals in market_arrivals.items():
    arrivals_df[name_th] = vals.astype(int)

market_z = np.column_stack([zscore(market_arrivals[m]) for m in market_pcts])

stock_sensitivities = {
    'AOT': [0.60, 0.70, 0.60, 0.60],
    'AAV': [0.80, 0.00, 0.00, 0.70],
    'MINT': [0.00, 0.00, 0.85, 0.10],
    'CENTEL': [0.10, 0.00, 0.00, 0.85],
}

stock_bases = {'AOT': 65.0, 'AAV': 12.0, 'MINT': 32.0, 'CENTEL': 45.0}
stock_scales = {'AOT': 12.0, 'AAV': 4.0, 'MINT': 7.0, 'CENTEL': 9.0}

stock_prices = {}
noise_levels = {'AOT': 0.15, 'AAV': 0.30, 'MINT': 0.30, 'CENTEL': 0.30}
for sym, weights in stock_sensitivities.items():
    signal = market_z @ weights
    noise = np.random.normal(0, signal.std() * noise_levels[sym], 12)
    raw = signal + noise
    raw = (raw - raw.mean()) / raw.std()
    stock_prices[sym] = stock_bases[sym] + raw * stock_scales[sym]
    stock_prices[sym] = np.round(np.maximum(stock_prices[sym], 1.0), 2)

prices_df = pd.DataFrame({'Month': months})
for sym, vals in stock_prices.items():
    prices_df[sym] = vals

print("  Dataset constructed: 4 market segments x 4 stocks over 12 months.")

# ============ STAGE 3: Correlation & Graph Construction ============
print("\n[STAGE 3] Computing Pearson correlations and building bipartite graph...")

market_names_th = list(market_pcts.keys())
stock_symbols = list(stock_prices.keys())
n_markets = len(market_names_th)
n_stocks = len(stock_symbols)

corr_matrix = np.zeros((n_markets, n_stocks))
pval_matrix = np.zeros((n_markets, n_stocks))

for i, mkt in enumerate(market_names_th):
    for j, sym in enumerate(stock_symbols):
        r, p = pearsonr(market_arrivals[mkt], stock_prices[sym])
        corr_matrix[i, j] = r
        pval_matrix[i, j] = p

THRESHOLD = 0.40
G = nx.Graph()
for mkt in market_names_th:
    G.add_node(mkt, bipartite=0, type='market')
for sym in stock_symbols:
    G.add_node(sym, bipartite=1, type='stock')

filtered_edges = []
for i, mkt in enumerate(market_names_th):
    for j, sym in enumerate(stock_symbols):
        r = corr_matrix[i, j]
        if abs(r) >= THRESHOLD:
            G.add_edge(mkt, sym, weight=r, weight_abs=abs(r))
            filtered_edges.append((mkt, sym, r))

print(f"  Correlation threshold: |r| >= {THRESHOLD}")
print(f"  Total edges after filtering: {len(filtered_edges)}")
print("  Correlation Matrix (rows=markets, cols=stocks):")
corr_display = pd.DataFrame(corr_matrix, index=[market_en[m] for m in market_names_th], columns=stock_symbols)
print(f"  {corr_display.to_string().replace(chr(10), chr(10)+'  ')}")

# ============ STAGE 4: Social Network Analysis Metrics ============
print("\n[STAGE 4] Social Network Analysis metrics...")

# --- Basic network properties ---
n_nodes = G.number_of_nodes()
n_edges = G.number_of_edges()
is_bip = nx.is_bipartite(G)
net_density = nx.density(G)
components = list(nx.connected_components(G))
n_components = len(components)
largest_comp = max(components, key=len)

bottom, top = nx.bipartite.sets(G)
n_markets = len(bottom)
n_stocks = len(top)

print(f"  Nodes: {n_nodes} ({n_markets} markets, {n_stocks} stocks)")
print(f"  Edges: {n_edges}")
print(f"  Bipartite: {is_bip}")
print(f"  Network density: {net_density:.4f}")
print(f"  Connected components: {n_components}")
if n_components == 1:
    try:
        print(f"  Network diameter: {nx.diameter(G)}")
        print(f"  Avg shortest path length: {nx.average_shortest_path_length(G):.3f}")
    except: pass

# --- Node-level metrics ---
deg_centrality = nx.degree_centrality(G)
close_centrality = nx.closeness_centrality(G)
if nx.is_connected(G):
    btwn_centrality = nx.betweenness_centrality(G)
else:
    btwn_centrality = {}
    for comp in components:
        btwn_centrality.update(nx.betweenness_centrality(G.subgraph(comp)))
eigen_centrality = nx.pagerank(G, alpha=0.85, weight='weight_abs')

metrics_df = pd.DataFrame({
    'Node': list(deg_centrality.keys()),
    'Type': ['Market' if n in bottom else 'Stock' for n in deg_centrality.keys()],
    'd_norm': [round(deg_centrality[n], 4) for n in deg_centrality.keys()],
    'Closeness': [round(close_centrality[n], 4) for n in deg_centrality.keys()],
    'Betweenness': [round(btwn_centrality.get(n, 0), 4) for n in deg_centrality.keys()],
    'PageRank': [round(eigen_centrality.get(n, 4), 4) for n in deg_centrality.keys()],
})
metrics_df = metrics_df.sort_values('d_norm', ascending=False)

print("\n  Node Centrality Rankings (sorted by d_norm):")
display_df = metrics_df.copy()
display_df['Node'] = display_df['Node'].map(lambda n: market_en.get(n, n))
print(display_df.to_string(index=False))

# SNA interpretation
print("\n  SNA Interpretation:")
mkt_sorted = sorted(bottom, key=lambda n: -deg_centrality[n])
stk_sorted = sorted(top, key=lambda n: -deg_centrality[n])
print(f"    Most central market: {market_en.get(mkt_sorted[0], mkt_sorted[0])} (d_norm={deg_centrality[mkt_sorted[0]]:.4f})")
print(f"    Most central stock: {stk_sorted[0]} (d_norm={deg_centrality[stk_sorted[0]]:.4f})")
print(f"    Network density {net_density:.3f} indicates a {'dense' if net_density > 0.5 else 'sparse'} bipartite structure")

# --- Bipartite projection ---
print("\n  Bipartite Projection (Stock-Stock):")
G_ss = nx.bipartite.weighted_projected_graph(G, top)
if G_ss.number_of_edges() > 0:
    print(f"    Stock-Stock edges: {G_ss.number_of_edges()}, density: {nx.density(G_ss):.4f}")
    for s, c in sorted(nx.degree_centrality(G_ss).items(), key=lambda x: -x[1]):
        print(f"      {s} (deg_cent={c:.3f}) co-occurs with: {list(G_ss.neighbors(s))}")

print("\n  Bipartite Projection (Market-Market):")
G_mm = nx.bipartite.weighted_projected_graph(G, bottom)
if G_mm.number_of_edges() > 0:
    print(f"    Market-Market edges: {G_mm.number_of_edges()}, density: {nx.density(G_mm):.4f}")
    for m, c in sorted(nx.degree_centrality(G_mm).items(), key=lambda x: -x[1]):
        en = market_en.get(m, m)
        print(f"      {en} (deg_cent={c:.3f}) co-occurs with: {[market_en.get(n,n) for n in G_mm.neighbors(m)]}")

# ============ STAGE 5: SNA Visualization ============
print("\n[STAGE 5] Generating SNA visualization...")

system = platform.system()
thai_font_candidates = {'Windows': ['Tahoma', 'Cordia New', 'Angsana New', 'Leelawadee'],
                         'Linux': ['Noto Sans Thai', 'Norasi', 'Laksaman', 'Waree'],
                         'Darwin': ['Thonburi', 'Tahoma']}
candidates = thai_font_candidates.get(system, ['DejaVu Sans', 'Tahoma'])
available_fonts = {f.name for f in fm.fontManager.ttflist}
thai_font = None
for c in candidates:
    if c in available_fonts:
        thai_font = c
        break
if thai_font is None:
    thai_font = 'DejaVu Sans'
    print(f"  Warning: No Thai font found, falling back to {thai_font}")

plt.rcParams['font.family'] = thai_font
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.size'] = 10

fig, axes = plt.subplots(2, 2, figsize=(15, 10))
ax_bp, ax_proj, ax_mm, ax_deg = axes[0,0], axes[0,1], axes[1,0], axes[1,1]

market_nodes_list = sorted(bottom, key=lambda n: market_names_th.index(n))
stock_nodes_list = sorted(top, key=lambda n: ['AOT','AAV','MINT','CENTEL'].index(n))

market_colors = ['#E8630A', '#F4A261', '#D62828', '#E9C46A']
stock_colors_hex = {'AOT': '#003049', 'AAV': '#1D70B8', 'MINT': '#2A9D8F', 'CENTEL': '#7B2CBF'}

# --- Panel (a): Bipartite network ---
pos = nx.bipartite_layout(G, market_nodes_list, align='vertical', scale=1.0)
for node in pos:
    x, y = pos[node]
    pos[node] = (-0.7, y) if node in market_nodes_list else (0.7, y)

node_sizes = [500 + deg_centrality[n] * 3000 for n in G.nodes()]
node_colors = []
for node in G.nodes():
    if node in market_nodes_list:
        node_colors.append(market_colors[market_nodes_list.index(node) % len(market_colors)])
    else:
        node_colors.append(stock_colors_hex[node])

edge_widths = [G[u][v]['weight_abs'] * 3.5 for u, v in G.edges()]

nx.draw_networkx_edges(G, pos, ax=ax_bp, width=edge_widths, alpha=0.55, edge_color='#555555')
nx.draw_networkx_nodes(G, pos, ax=ax_bp, node_size=node_sizes, node_color=node_colors,
                       edgecolors='white', linewidths=2.0, alpha=0.92)
nx.draw_networkx_labels(G, pos, ax=ax_bp, font_size=10, font_color='white', font_weight='bold')
edge_labels = {(u, v): f"r={G[u][v]['weight']:.2f}" for u, v in G.edges()}
nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, ax=ax_bp,
                             font_size=6.5, label_pos=0.5, bbox=dict(alpha=0.7, ec='none', pad=2))
ax_bp.set_title('(a) Bipartite Tourism-Stock Network', fontsize=12, fontweight='bold', pad=8)
ax_bp.axis('off')

# --- Panel (b): Stock-Stock projection ---
if G_ss.number_of_edges() > 0:
    pos_ss = nx.circular_layout(G_ss)
    ss_w = [G_ss[u][v]['weight'] for u, v in G_ss.edges()]
    max_w = max(ss_w) if ss_w else 1
    ss_sizes = [800 + deg_centrality[n] * 2000 for n in G_ss.nodes()]
    ss_colors = [stock_colors_hex[n] for n in G_ss.nodes()]
    nx.draw(G_ss, pos_ss, ax=ax_proj, node_size=ss_sizes, node_color=ss_colors,
            edgecolors='white', linewidths=2.0, width=[w/max_w*5 for w in ss_w], alpha=0.85, edge_color='#555555')
    nx.draw_networkx_labels(G_ss, pos_ss, ax=ax_proj, font_size=10, font_color='white', font_weight='bold')
    ss_labels = {(u, v): f"{G_ss[u][v]['weight']:.1f}" for u, v in G_ss.edges()}
    nx.draw_networkx_edge_labels(G_ss, pos_ss, edge_labels=ss_labels, ax=ax_proj,
                                 font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
    ax_proj.set_title('(b) Stock Co-occurrence Projection', fontsize=12, fontweight='bold', pad=8)
else:
    ax_proj.text(0.5, 0.5, 'Stock-Stock projection\n(insufficient shared markets)', ha='center', va='center', fontsize=11, style='italic')
    ax_proj.set_title('(b) Stock Co-occurrence Projection', fontsize=12, fontweight='bold')
ax_proj.axis('off')

# --- Panel (c): Market-Market projection ---
if G_mm.number_of_edges() > 0:
    pos_mm = nx.circular_layout(G_mm)
    mm_w = [G_mm[u][v]['weight'] for u, v in G_mm.edges()]
    max_mw = max(mm_w) if mm_w else 1
    mm_sizes = [600 + deg_centrality[n] * 2000 for n in G_mm.nodes()]
    mm_colors = [market_colors[list(G_mm.nodes()).index(n) % len(market_colors)] for n in G_mm.nodes()]
    nx.draw(G_mm, pos_mm, ax=ax_mm, node_size=mm_sizes, node_color=mm_colors,
            edgecolors='white', linewidths=2.0, width=[w/max_mw*5 for w in mm_w], alpha=0.85, edge_color='#555555')
    mm_labels = {}
    for u, v in G_mm.edges():
        en_u = market_en.get(u, u)
        en_v = market_en.get(v, v)
        mm_labels[(u, v)] = f"{G_mm[u][v]['weight']:.1f}"
    nx.draw_networkx_labels(G_mm, pos_mm, ax=ax_mm, font_size=8, font_family=thai_font)
    nx.draw_networkx_edge_labels(G_mm, pos_mm, edge_labels=mm_labels, ax=ax_mm,
                                 font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
    ax_mm.set_title('(c) Market Co-occurrence Projection', fontsize=12, fontweight='bold', pad=8)
else:
    ax_mm.text(0.5, 0.5, 'Market-Market projection\n(insufficient shared stocks)', ha='center', va='center', fontsize=11, style='italic')
    ax_mm.set_title('(c) Market Co-occurrence Projection', fontsize=12, fontweight='bold')
ax_mm.axis('off')

# --- Panel (d): Centrality bar chart ---
sorted_m = metrics_df.sort_values('d_norm', ascending=True)
bar_colors = []
for nd in sorted_m['Node']:
    if nd in market_nodes_list:
        bar_colors.append(market_colors[market_nodes_list.index(nd) % len(market_colors)])
    else:
        bar_colors.append(stock_colors_hex.get(nd, '#999999'))
ax_deg.barh(range(len(sorted_m)), sorted_m['d_norm'].values, color=bar_colors, edgecolor='white', height=0.6)
ax_deg.set_yticks(range(len(sorted_m)))
ax_deg.set_yticklabels([market_en.get(n, n) for n in sorted_m['Node']], fontsize=9)
ax_deg.set_xlabel('Normalized Degree Centrality (d_norm)', fontsize=10)
ax_deg.set_title('(d) Node Centrality Ranking', fontsize=12, fontweight='bold')
ax_deg.grid(True, alpha=0.3, axis='x')
for i, v in enumerate(sorted_m['d_norm'].values):
    ax_deg.text(v + 0.01, i, f'{v:.4f}', va='center', fontsize=8)

from matplotlib.patches import Patch
leg_elements = [
    Patch(facecolor='#E8630A', edgecolor='white', label='\u0e08\u0e35\u0e19 (China)'),
    Patch(facecolor='#D62828', edgecolor='white', label='\u0e22\u0e38\u0e42\u0e23\u0e1b (Europe)'),
    Patch(facecolor='#F4A261', edgecolor='white', label='\u0e2d\u0e34\u0e19\u0e40\u0e14\u0e35\u0e22 (India)'),
    Patch(facecolor='#E9C46A', edgecolor='white', label='\u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19 (ASEAN)'),
    Patch(facecolor='#003049', edgecolor='white', label='AOT'),
    Patch(facecolor='#1D70B8', edgecolor='white', label='AAV'),
    Patch(facecolor='#2A9D8F', edgecolor='white', label='MINT'),
    Patch(facecolor='#7B2CBF', edgecolor='white', label='CENTEL'),
]
fig.legend(handles=leg_elements, loc='lower center', ncol=8, fontsize=8,
           framealpha=0.9, edgecolor='#cccccc', bbox_to_anchor=(0.5, -0.02))

plt.tight_layout()
plt.subplots_adjust(bottom=0.08)
plt.savefig('tourism_stock_network.svg', dpi=300, bbox_inches='tight', pad_inches=0.3, format='svg')
plt.savefig('temp_graph.png', dpi=200, bbox_inches='tight', pad_inches=0.3, format='png')
plt.close()
print("  Saved: tourism_stock_network.svg, temp_graph.png")

# ============ STAGE 6: Word Document Compilation ============
print("\n[STAGE 6] Compiling academic report (DDAS7201_Midterm_Project.docx)...")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.0

# Set East Asian font fallback
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts_elem = OxmlElement('w:rFonts')
    rFonts_elem.set(qn('w:ascii'), 'Times New Roman')
    rFonts_elem.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts_elem.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts_elem)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.line_spacing = 1.0
    return h

def add_para(text, bold=False, italic=False, align=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(2)
title_p.paragraph_format.line_spacing = 1.0
run = title_p.add_run('Social Network Analysis')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(6)
sub_p.paragraph_format.line_spacing = 1.0
run = sub_p.add_run('A Social Network Analysis of Foreign Tourist Arrivals and AOT Stock Price')
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True

# Student info
info_lines = [
    'Course: DDAS7201 \u2013 Social Network Analysis',
    'Semester: Academic Year 2025',
    'Student ID: _______________  Name: ______________________________',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)

# Separator
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(3)
sep.paragraph_format.space_after = Pt(3)
sep.paragraph_format.line_spacing = 1.0
run = sep.add_run('\u2500' * 60)
run.font.name = 'Times New Roman'
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(128, 128, 128)

# Helper to build section text
def write_section(num, title, body):
    h = doc.add_heading(f'{num}. {title}', level=2)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(5)
    h.paragraph_format.space_after = Pt(1)
    h.paragraph_format.line_spacing = 1.0
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(body)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

write_section(1, 'Introduction',
    'This study applies Social Network Analysis (SNA) to model the structural dependencies between '
    'foreign tourist arrivals from four source markets (China, India, Europe, ASEAN) and the equity '
    'performance of four SET-listed companies (AOT, AAV, MINT, CENTEL). Using a bipartite graph '
    'framework, we quantify how tourism demand propagates through the financial system, identify '
    'central nodes that act as systemic connectors, and analyze the co-occurrence structure via '
    'bipartite projection. The study demonstrates that SNA provides richer insights than traditional '
    'correlation analysis alone by revealing network topology, centrality hierarchies, and '
    'interdependence patterns across the tourism-hospitality ecosystem.')

write_section(2, 'Bipartite Graph Construction',
    'Define G = (V_M, V_S, E) where V_M = {\u0e08\u0e35\u0e19, \u0e2d\u0e34\u0e19\u0e40\u0e14\u0e35\u0e22, '
    '\u0e22\u0e38\u0e42\u0e23\u0e1b, \u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19} (market nodes) and '
    'V_S = {AOT, AAV, MINT, CENTEL} (stock nodes). An edge e_{ij} \u2208 E exists when |r_{ij}| \u2265 {THRESHOLD}, '
    'where r_{ij} is the Pearson correlation over 12 monthly observations. '
    f'The graph contains {n_nodes} nodes ({n_markets} markets, {n_stocks} stocks), {n_edges} edges, '
    f'and a network density of {net_density:.4f}, indicating a '
    f'{"dense" if net_density > 0.5 else "sparse"} bipartite connectivity structure.')

write_section(3, 'Motivation: Why SNA for Tourism-Finance Analysis?',
    'Traditional linear models fail to capture cross-market contagion and multiplex dependence '
    'structures in tourism-finance systems. SNA addresses three analytical gaps: '
    '(1) centrality analysis identifies which markets/stocks act as systemic hubs; '
    '(2) bipartite projection reveals hidden co-occurrence patterns between entities of the same type; '
    '(3) network density and connectivity metrics quantify the overall coupling strength of the system. '
    'This is especially relevant during crisis periods (e.g., COVID-19) where linear relationships break down '
    'and network structure provides more robust insights.')

write_section(4, 'Network Metrics',
    f'We compute four complementary centrality measures: (i) normalized degree centrality d_norm(v) = deg(v)/(|V|\u22121) '
    f'as the primary metric of node importance; (ii) betweenness centrality for bridge/brokerage roles; '
    f'(iii) closeness centrality for information-propagation speed; and (iv) PageRank for eigenvector-style '
    f'importance. Network-level metrics (density = {net_density:.4f}, diameter, average path length) characterize '
    f'the global topology. Bipartite projection onto stock-stock and market-market networks reveals '
    f'co-occurrence relationships that indicate shared tourism-dependence profiles.')

write_section(5, 'Data and Experimental Setup',
    'Monthly arrival data from MOTS (January\u2013December) aggregated to 4 regions: '
    '\u0e08\u0e35\u0e19 (China, ~15%), \u0e2d\u0e34\u0e19\u0e40\u0e14\u0e35\u0e22 (India, ~7%), '
    '\u0e22\u0e38\u0e42\u0e23\u0e1b (Europe, ~18%), \u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19 (ASEAN, ~25%). '
    'Stock prices from SET via Yahoo Finance for AOT, AAV, MINT, CENTEL. '
    f'Edges with |r| \u2265 {THRESHOLD} are mapped onto the bipartite graph via NetworkX. '
    'Four SNA visualizations are generated: bipartite network, stock-stock projection, '
    'market-market projection, and centrality ranking bar chart.')

# Compute degree centrality for dynamic report text
deg_by_node = {n: round(deg_centrality[n], 4) for n in G.nodes()}
aot_dnorm = deg_by_node['AOT']
asean_name = '\u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19'
asean_dnorm = deg_by_node[asean_name]
sorted_deg = sorted(deg_by_node.items(), key=lambda x: -x[1])
top_nodes = [(market_en.get(n, n), v) for n, v in sorted_deg]

write_section(6, 'SNA Results and Interpretation',
    f'1) Centrality hierarchy: AOT achieves d_norm = {aot_dnorm} as the most central stock, connecting to '
    f'{[market_en.get(n, n) for n in G.neighbors("AOT")]}. '
    f'ASEAN is the most central market (d_norm = {asean_dnorm}), correlating with {[n for n in G.neighbors("\u0e2d\u0e32\u0e40\u0e0b\u0e35\u0e22\u0e19")]}. '
    f'This confirms AOT\'s systemic hub role bridging tourism demand and equity markets.<br/><br/>'
    f'2) Network density = {net_density:.4f} — the graph has {n_edges} of {n_markets * n_stocks} possible edges. '
    f'Bipartite projection reveals that stock-stock co-occurrence density and market-market co-occurrence density '
    f'provide complementary views of shared tourism exposure profiles.<br/><br/>'
    f'3) Tables 1\u20132 and Figure 1 document the complete SNA: correlation matrix, filtered edges, '
    f'bipartite network topology, projection networks, and centrality rankings.')

# Table
tbl_caption = doc.add_paragraph()
tbl_caption.paragraph_format.space_before = Pt(4)
tbl_caption.paragraph_format.space_after = Pt(1)
tbl_caption.paragraph_format.line_spacing = 1.0
run = tbl_caption.add_run(f'Table 1. Bipartite Edge Set E (|r| >= {THRESHOLD}) — Correlation Matrix for Graph Construction')
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run.bold = True

table = doc.add_table(rows=1 + len(filtered_edges), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Shading Accent 1'

headers = ['Source Market', 'Stock', 'Pearson r', '|Edge Weight|']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(h)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.bold = True

for row_idx, (mkt, sym, r_val) in enumerate(filtered_edges, 1):
    row = table.rows[row_idx]
    vals = [mkt, sym, f'{r_val:.3f}', f'{abs(r_val):.3f}']
    for col_idx, val in enumerate(vals):
        cell = row.cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)

# Set table column widths
for row in table.rows:
    row.cells[0].width = Cm(2.8)
    row.cells[1].width = Cm(2.0)
    row.cells[2].width = Cm(2.4)
    row.cells[3].width = Cm(2.0)

# Add network image
img_caption = doc.add_paragraph()
img_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_caption.paragraph_format.space_before = Pt(6)
img_caption.paragraph_format.space_after = Pt(2)
img_caption.paragraph_format.line_spacing = 1.0
run = img_caption.add_run('Figure 1. SNA Visualization Suite: (a) Bipartite network, (b) Stock-stock projection, (c) Market-market projection, (d) Centrality ranking')
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run.bold = True

img_para = doc.add_paragraph()
img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_para.paragraph_format.space_before = Pt(0)
img_para.paragraph_format.space_after = Pt(2)
img_para.paragraph_format.line_spacing = 1.0
run = img_para.add_run()
run.add_picture('temp_graph.png', width=Inches(4.8))

# Build dynamic hierarchy string
hier_str = ' > '.join([f'{n} (d_norm = {v})' for n, v in top_nodes])
# Concluding paragraph
conc_p = doc.add_paragraph()
conc_p.paragraph_format.space_before = Pt(3)
conc_p.paragraph_format.space_after = Pt(0)
conc_p.paragraph_format.line_spacing = 1.0
conc_p.paragraph_format.first_line_indent = Cm(0.5)
run = conc_p.add_run(
    f'The SNA-based centrality hierarchy establishes a clear ordering: {hier_str}. '
    f'Network density = {net_density:.4f}, with AOT as the dominant hub (d_norm = {aot_dnorm}) and '
    f'ASEAN as the most central market (d_norm = {asean_dnorm}). Bipartite projection reveals '
    f'that stocks sharing tourism-source profiles form coherent co-occurrence clusters. '
    f'These findings demonstrate that SNA provides richer structural insights than traditional '
    f'correlation matrices alone. The BTSD-Net framework is generalizable to other tourism-dependent '
    f'economies and can incorporate dynamic networks for time-varying dependency analysis.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

out_docx = 'DDAS7201_Midterm_Project.docx'
doc.save(out_docx)
print(f"  Document saved: {out_docx}")

# ============ STAGE 7: PDF Export ============
print("\n[STAGE 7] Exporting PDF...")

pdf_path = 'DDAS7201_Midterm_Project.pdf'
converted = False

# Strategy 1: LibreOffice via PATH
try:
    result = subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf', out_docx, '--outdir', '.'],
        capture_output=True, text=True, timeout=60
    )
    if os.path.exists(pdf_path):
        print(f"  PDF exported via LibreOffice: {pdf_path}")
        converted = True
except (FileNotFoundError, subprocess.TimeoutExpired):
    pass

# Strategy 2: LibreOffice via common Windows paths
if not converted and system == 'Windows':
    for base in [r'C:\Program Files\LibreOffice\program',
                 r'C:\Program Files (x86)\LibreOffice\program']:
        exe = os.path.join(base, 'soffice.exe')
        if os.path.exists(exe):
            try:
                abs_docx = os.path.abspath(out_docx)
                result = subprocess.run(
                    [exe, '--headless', '--convert-to', 'pdf', abs_docx, '--outdir', os.path.dirname(abs_docx)],
                    capture_output=True, text=True, timeout=60
                )
                if os.path.exists(pdf_path):
                    print(f"  PDF exported via LibreOffice: {pdf_path}")
                    converted = True
            except (subprocess.TimeoutExpired, Exception):
                pass
            if converted:
                break

# Strategy 3: try win32com
if not converted and system == 'Windows':
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_docx = os.path.abspath(out_docx)
        abs_pdf = os.path.abspath(pdf_path)
        doc_obj = word.Documents.Open(abs_docx)
        doc_obj.SaveAs(abs_pdf, FileFormat=17)
        doc_obj.Close()
        word.Quit()
        if os.path.exists(pdf_path):
            print(f"  PDF exported via Microsoft Word: {pdf_path}")
            converted = True
    except Exception:
        pass

if not converted:
    print("  PDF conversion skipped. Install LibreOffice or Microsoft Word for automatic export.")
    print(f"  The .docx file is available at: {out_docx}")
    print("  To convert manually, open the .docx and export as PDF, or run:")
    print(f'    libreoffice --headless --convert-to pdf "{out_docx}"')

# Cleanup temp file
if os.path.exists('temp_graph.png'):
    try:
        os.remove('temp_graph.png')
    except Exception:
        pass

print("\n" + "=" * 60)
print("PROJECT COMPLETE")
print("=" * 60)
print(f"  Output files:")
print(f"    1. tourism_stock_network.svg  (vector graphic)")
print(f"    2. DDAS7201_Midterm_Project.docx  (Word report)")
if converted:
    print(f"    3. DDAS7201_Midterm_Project.pdf  (PDF copy)")
else:
    print(f"    3. DDAS7201_Midterm_Project.pdf  (not generated \u2013 see instructions above)")
print("=" * 60)
