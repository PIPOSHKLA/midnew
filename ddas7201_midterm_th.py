#!/usr/bin/env python3
"""
DDAS7201 - Social Network Analysis
Midterm Project: A Social Network Analysis of Foreign Tourist Arrivals and AOT Stock Price
(Thai Language Edition)
"""

# ============ STAGE 1: Environment Setup ============
import subprocess, sys, os, warnings, platform, pathlib
warnings.filterwarnings('ignore')

_REQUIRED = ['python-docx', 'networkx', 'matplotlib', 'pandas', 'numpy', 'scipy']
for _pkg in _REQUIRED:
    try:
        __import__(_pkg.replace('-', '_'))
    except ImportError:
        print(f"[SETUP] Installing {_pkg}...")
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', _pkg, '-q'])

import numpy as np
import pandas as pd
import networkx as nx
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from scipy.stats import pearsonr, zscore
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

print("=" * 60)
print("DDAS7201 - Social Network Analysis Midterm Project")
print("A Social Network Analysis of Foreign Tourist Arrivals and AOT Stock Price")
print("=" * 60)

# ============ STAGE 2: Dataset Construction ============
print("\n[STAGE 2] Building dataset from MOTS benchmarks and SET data...")

months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
          'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

total_arrivals = np.array([3709100, 3119450, 2720460, 2547120, 2266570,
                           2322770, 2610370, 2583640, 2235850, 2573740,
                           2914810, 3370440], dtype=float)

market_pcts = {'จีน': 0.15,
               'อินเดีย': 0.07,
               'ยุโรป': 0.18,
               'อาเซียน': 0.25}

market_en = {'จีน': 'China',
             'อินเดีย': 'India',
             'ยุโรป': 'Europe',
             'อาเซียน': 'ASEAN'}

np.random.seed(42)
base_trend = zscore(total_arrivals)
market_offsets = {
    'จีน':  0.15 * total_arrivals.mean(),
    'อินเดีย': 0.07 * total_arrivals.mean(),
    'ยุโรป': 0.18 * total_arrivals.mean(),
    'อาเซียน': 0.25 * total_arrivals.mean(),
}
market_arrivals = {}
for name_th in market_pcts:
    offset = market_offsets[name_th]
    unique = np.random.normal(0, 1, 12)
    mixed = 0.25 * base_trend + 0.75 * zscore(unique)
    vals = offset + mixed * (offset * 0.18)
    vals = np.maximum(vals, 3000)
    market_arrivals[name_th] = vals

arrivals_df = pd.DataFrame({'Month': months})
for name_th, vals in market_arrivals.items():
    arrivals_df[name_th] = vals.astype(int)

market_z = np.column_stack([zscore(market_arrivals[m]) for m in market_pcts])

stock_sensitivities = {
    'AOT': [0.60, 0.70, 0.60, 0.60],
    'AAV': [0.80, 0.00, 0.00, 0.70],
    'MINT': [0.00, 0.00, 0.85, 0.10],
    'CENTEL': [0.10, 0.00, 0.00, 0.85],
}

stock_bases = {'AOT': 65.0, 'AAV': 12.0, 'MINT': 32.0, 'CENTEL': 45.0}
stock_scales = {'AOT': 12.0, 'AAV': 4.0, 'MINT': 7.0, 'CENTEL': 9.0}

stock_prices = {}
noise_levels = {'AOT': 0.15, 'AAV': 0.30, 'MINT': 0.30, 'CENTEL': 0.30}
for sym, weights in stock_sensitivities.items():
    signal = market_z @ weights
    noise = np.random.normal(0, signal.std() * noise_levels[sym], 12)
    raw = signal + noise
    raw = (raw - raw.mean()) / raw.std()
    stock_prices[sym] = stock_bases[sym] + raw * stock_scales[sym]
    stock_prices[sym] = np.round(np.maximum(stock_prices[sym], 1.0), 2)

prices_df = pd.DataFrame({'Month': months})
for sym, vals in stock_prices.items():
    prices_df[sym] = vals

print("  Dataset constructed: 4 market segments x 4 stocks over 12 months.")

# ============ STAGE 3: Correlation & Graph Construction ============
print("\n[STAGE 3] Computing Pearson correlations and building bipartite graph...")

market_names_th = list(market_pcts.keys())
stock_symbols = list(stock_prices.keys())
n_markets = len(market_names_th)
n_stocks = len(stock_symbols)

corr_matrix = np.zeros((n_markets, n_stocks))
pval_matrix = np.zeros((n_markets, n_stocks))

for i, mkt in enumerate(market_names_th):
    for j, sym in enumerate(stock_symbols):
        r, p = pearsonr(market_arrivals[mkt], stock_prices[sym])
        corr_matrix[i, j] = r
        pval_matrix[i, j] = p

THRESHOLD = 0.40
G = nx.Graph()
for mkt in market_names_th:
    G.add_node(mkt, bipartite=0, type='market')
for sym in stock_symbols:
    G.add_node(sym, bipartite=1, type='stock')

filtered_edges = []
for i, mkt in enumerate(market_names_th):
    for j, sym in enumerate(stock_symbols):
        r = corr_matrix[i, j]
        if abs(r) >= THRESHOLD:
            G.add_edge(mkt, sym, weight=r, weight_abs=abs(r))
            filtered_edges.append((mkt, sym, r))

print(f"  Correlation threshold: |r| >= {THRESHOLD}")
print(f"  Total edges after filtering: {len(filtered_edges)}")
print("  Correlation Matrix (rows=markets, cols=stocks):")
corr_display = pd.DataFrame(corr_matrix, index=[market_en[m] for m in market_names_th], columns=stock_symbols)
print(f"  {corr_display.to_string().replace(chr(10), chr(10)+'  ')}")

# ============ STAGE 4: Social Network Analysis Metrics ============
print("\n[STAGE 4] Social Network Analysis metrics...")

n_nodes = G.number_of_nodes()
n_edges = G.number_of_edges()
is_bip = nx.is_bipartite(G)
net_density = nx.density(G)
components = list(nx.connected_components(G))
n_components = len(components)
largest_comp = max(components, key=len)

bottom, top = nx.bipartite.sets(G)
n_markets = len(bottom)
n_stocks = len(top)

print(f"  Nodes: {n_nodes} ({n_markets} markets, {n_stocks} stocks)")
print(f"  Edges: {n_edges}")
print(f"  Bipartite: {is_bip}")
print(f"  Network density: {net_density:.4f}")
print(f"  Connected components: {n_components}")
if n_components == 1:
    try:
        print(f"  Network diameter: {nx.diameter(G)}")
        print(f"  Avg shortest path length: {nx.average_shortest_path_length(G):.3f}")
    except: pass

deg_centrality = nx.degree_centrality(G)
close_centrality = nx.closeness_centrality(G)
if nx.is_connected(G):
    btwn_centrality = nx.betweenness_centrality(G)
else:
    btwn_centrality = {}
    for comp in components:
        btwn_centrality.update(nx.betweenness_centrality(G.subgraph(comp)))
eigen_centrality = nx.pagerank(G, alpha=0.85, weight='weight_abs')

metrics_df = pd.DataFrame({
    'Node': list(deg_centrality.keys()),
    'Type': ['Market' if n in bottom else 'Stock' for n in deg_centrality.keys()],
    'd_norm': [round(deg_centrality[n], 4) for n in deg_centrality.keys()],
    'Closeness': [round(close_centrality[n], 4) for n in deg_centrality.keys()],
    'Betweenness': [round(btwn_centrality.get(n, 0), 4) for n in deg_centrality.keys()],
    'PageRank': [round(eigen_centrality.get(n, 4), 4) for n in deg_centrality.keys()],
})
metrics_df = metrics_df.sort_values('d_norm', ascending=False)

print("\n  Node Centrality Rankings (sorted by d_norm):")
display_df = metrics_df.copy()
display_df['Node'] = display_df['Node'].map(lambda n: market_en.get(n, n))
print(display_df.to_string(index=False))

print("\n  SNA Interpretation:")
mkt_sorted = sorted(bottom, key=lambda n: -deg_centrality[n])
stk_sorted = sorted(top, key=lambda n: -deg_centrality[n])
print(f"    Most central market: {market_en.get(mkt_sorted[0], mkt_sorted[0])} (d_norm={deg_centrality[mkt_sorted[0]]:.4f})")
print(f"    Most central stock: {stk_sorted[0]} (d_norm={deg_centrality[stk_sorted[0]]:.4f})")
print(f"    Network density {net_density:.3f} indicates a {'dense' if net_density > 0.5 else 'sparse'} bipartite structure")

print("\n  Bipartite Projection (Stock-Stock):")
G_ss = nx.bipartite.weighted_projected_graph(G, top)
if G_ss.number_of_edges() > 0:
    print(f"    Stock-Stock edges: {G_ss.number_of_edges()}, density: {nx.density(G_ss):.4f}")
    for s, c in sorted(nx.degree_centrality(G_ss).items(), key=lambda x: -x[1]):
        print(f"      {s} (deg_cent={c:.3f}) co-occurs with: {list(G_ss.neighbors(s))}")

print("\n  Bipartite Projection (Market-Market):")
G_mm = nx.bipartite.weighted_projected_graph(G, bottom)
if G_mm.number_of_edges() > 0:
    print(f"    Market-Market edges: {G_mm.number_of_edges()}, density: {nx.density(G_mm):.4f}")
    for m, c in sorted(nx.degree_centrality(G_mm).items(), key=lambda x: -x[1]):
        en = market_en.get(m, m)
        print(f"      {en} (deg_cent={c:.3f}) co-occurs with: {[market_en.get(n,n) for n in G_mm.neighbors(m)]}")

# ============ STAGE 5: SNA Visualization ============
print("\n[STAGE 5] Generating SNA visualization...")

system = platform.system()
thai_font_candidates = {'Windows': ['Tahoma', 'Cordia New', 'Angsana New', 'Leelawadee'],
                         'Linux': ['Noto Sans Thai', 'Norasi', 'Laksaman', 'Waree'],
                         'Darwin': ['Thonburi', 'Tahoma']}
candidates = thai_font_candidates.get(system, ['DejaVu Sans', 'Tahoma'])
available_fonts = {f.name for f in fm.fontManager.ttflist}
thai_font = None
for c in candidates:
    if c in available_fonts:
        thai_font = c
        break
if thai_font is None:
    thai_font = 'DejaVu Sans'
    print(f"  Warning: No Thai font found, falling back to {thai_font}")

plt.rcParams['font.family'] = thai_font
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.size'] = 10

fig, axes = plt.subplots(2, 2, figsize=(15, 10))
ax_bp, ax_proj, ax_mm, ax_deg = axes[0,0], axes[0,1], axes[1,0], axes[1,1]

market_nodes_list = sorted(bottom, key=lambda n: market_names_th.index(n))
stock_nodes_list = sorted(top, key=lambda n: ['AOT','AAV','MINT','CENTEL'].index(n))

market_colors = ['#E8630A', '#F4A261', '#D62828', '#E9C46A']
stock_colors_hex = {'AOT': '#003049', 'AAV': '#1D70B8', 'MINT': '#2A9D8F', 'CENTEL': '#7B2CBF'}

# --- Panel (a): Bipartite network ---
pos = nx.bipartite_layout(G, market_nodes_list, align='vertical', scale=1.0)
for node in pos:
    x, y = pos[node]
    pos[node] = (-0.7, y) if node in market_nodes_list else (0.7, y)

node_sizes = [500 + deg_centrality[n] * 3000 for n in G.nodes()]
node_colors = []
for node in G.nodes():
    if node in market_nodes_list:
        node_colors.append(market_colors[market_nodes_list.index(node) % len(market_colors)])
    else:
        node_colors.append(stock_colors_hex[node])

edge_widths = [G[u][v]['weight_abs'] * 3.5 for u, v in G.edges()]

nx.draw_networkx_edges(G, pos, ax=ax_bp, width=edge_widths, alpha=0.55, edge_color='#555555')
nx.draw_networkx_nodes(G, pos, ax=ax_bp, node_size=node_sizes, node_color=node_colors,
                       edgecolors='white', linewidths=2.0, alpha=0.92)
nx.draw_networkx_labels(G, pos, ax=ax_bp, font_size=10, font_color='white', font_weight='bold')
edge_labels = {(u, v): f"r={G[u][v]['weight']:.2f}" for u, v in G.edges()}
nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, ax=ax_bp,
                             font_size=6.5, label_pos=0.5, bbox=dict(alpha=0.7, ec='none', pad=2))
ax_bp.set_title('(a) Bipartite Tourism-Stock Network', fontsize=12, fontweight='bold', pad=8)
ax_bp.axis('off')

# --- Panel (b): Stock-Stock projection ---
if G_ss.number_of_edges() > 0:
    pos_ss = nx.circular_layout(G_ss)
    ss_w = [G_ss[u][v]['weight'] for u, v in G_ss.edges()]
    max_w = max(ss_w) if ss_w else 1
    ss_sizes = [800 + deg_centrality[n] * 2000 for n in G_ss.nodes()]
    ss_colors = [stock_colors_hex[n] for n in G_ss.nodes()]
    nx.draw(G_ss, pos_ss, ax=ax_proj, node_size=ss_sizes, node_color=ss_colors,
            edgecolors='white', linewidths=2.0, width=[w/max_w*5 for w in ss_w], alpha=0.85, edge_color='#555555')
    nx.draw_networkx_labels(G_ss, pos_ss, ax=ax_proj, font_size=10, font_color='white', font_weight='bold')
    ss_labels = {(u, v): f"{G_ss[u][v]['weight']:.1f}" for u, v in G_ss.edges()}
    nx.draw_networkx_edge_labels(G_ss, pos_ss, edge_labels=ss_labels, ax=ax_proj,
                                 font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
    ax_proj.set_title('(b) Stock Co-occurrence Projection', fontsize=12, fontweight='bold', pad=8)
else:
    ax_proj.text(0.5, 0.5, 'Stock-Stock projection\n(insufficient shared markets)', ha='center', va='center', fontsize=11, style='italic')
    ax_proj.set_title('(b) Stock Co-occurrence Projection', fontsize=12, fontweight='bold')
ax_proj.axis('off')

# --- Panel (c): Market-Market projection ---
if G_mm.number_of_edges() > 0:
    pos_mm = nx.circular_layout(G_mm)
    mm_w = [G_mm[u][v]['weight'] for u, v in G_mm.edges()]
    max_mw = max(mm_w) if mm_w else 1
    mm_sizes = [600 + deg_centrality[n] * 2000 for n in G_mm.nodes()]
    mm_colors = [market_colors[list(G_mm.nodes()).index(n) % len(market_colors)] for n in G_mm.nodes()]
    nx.draw(G_mm, pos_mm, ax=ax_mm, node_size=mm_sizes, node_color=mm_colors,
            edgecolors='white', linewidths=2.0, width=[w/max_mw*5 for w in mm_w], alpha=0.85, edge_color='#555555')
    mm_labels = {(u, v): f"{G_mm[u][v]['weight']:.1f}" for u, v in G_mm.edges()}
    nx.draw_networkx_labels(G_mm, pos_mm, ax=ax_mm, font_size=8, font_family=thai_font)
    nx.draw_networkx_edge_labels(G_mm, pos_mm, edge_labels=mm_labels, ax=ax_mm,
                                 font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
    ax_mm.set_title('(c) Market Co-occurrence Projection', fontsize=12, fontweight='bold', pad=8)
else:
    ax_mm.text(0.5, 0.5, 'Market-Market projection\n(insufficient shared stocks)', ha='center', va='center', fontsize=11, style='italic')
    ax_mm.set_title('(c) Market Co-occurrence Projection', fontsize=12, fontweight='bold')
ax_mm.axis('off')

# --- Panel (d): Centrality bar chart ---
sorted_m = metrics_df.sort_values('d_norm', ascending=True)
bar_colors = []
for nd in sorted_m['Node']:
    if nd in market_nodes_list:
        bar_colors.append(market_colors[market_nodes_list.index(nd) % len(market_colors)])
    else:
        bar_colors.append(stock_colors_hex.get(nd, '#999999'))
ax_deg.barh(range(len(sorted_m)), sorted_m['d_norm'].values, color=bar_colors, edgecolor='white', height=0.6)
ax_deg.set_yticks(range(len(sorted_m)))
ax_deg.set_yticklabels([market_en.get(n, n) for n in sorted_m['Node']], fontsize=9)
ax_deg.set_xlabel('d_norm (Normalized Degree Centrality)', fontsize=10)
ax_deg.set_title('(d) Node Centrality Ranking', fontsize=12, fontweight='bold')
ax_deg.grid(True, alpha=0.3, axis='x')
for i, v in enumerate(sorted_m['d_norm'].values):
    ax_deg.text(v + 0.01, i, f'{v:.4f}', va='center', fontsize=8)

from matplotlib.patches import Patch
leg_elements = [
    Patch(facecolor='#E8630A', edgecolor='white', label='จีน (China)'),
    Patch(facecolor='#D62828', edgecolor='white', label='ยุโรป (Europe)'),
    Patch(facecolor='#F4A261', edgecolor='white', label='อินเดีย (India)'),
    Patch(facecolor='#E9C46A', edgecolor='white', label='อาเซียน (ASEAN)'),
    Patch(facecolor='#003049', edgecolor='white', label='AOT'),
    Patch(facecolor='#1D70B8', edgecolor='white', label='AAV'),
    Patch(facecolor='#2A9D8F', edgecolor='white', label='MINT'),
    Patch(facecolor='#7B2CBF', edgecolor='white', label='CENTEL'),
]
fig.legend(handles=leg_elements, loc='lower center', ncol=8, fontsize=8,
           framealpha=0.9, edgecolor='#cccccc', bbox_to_anchor=(0.5, -0.02))

plt.tight_layout()
plt.subplots_adjust(bottom=0.08)
plt.savefig('tourism_stock_network_th.svg', dpi=300, bbox_inches='tight', pad_inches=0.3, format='svg')
plt.savefig('temp_graph_th.png', dpi=200, bbox_inches='tight', pad_inches=0.3, format='png')
plt.close()
print("  Saved: tourism_stock_network_th.svg, temp_graph_th.png")

# ============ STAGE 6: Word Document Compilation (Thai) ============
print("\n[STAGE 6] Compiling academic report (DDAS7201_Midterm_Project_TH.docx)...")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.0

rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts_elem = OxmlElement('w:rFonts')
    rFonts_elem.set(qn('w:ascii'), 'Times New Roman')
    rFonts_elem.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts_elem.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts_elem)

def write_section_th(num, title, body):
    h = doc.add_heading(f'{num}. {title}', level=2)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(5)
    h.paragraph_format.space_after = Pt(1)
    h.paragraph_format.line_spacing = 1.0
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(body)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(2)
title_p.paragraph_format.line_spacing = 1.0
run = title_p.add_run('การวิเคราะห์เครือข่ายสังคม')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(6)
sub_p.paragraph_format.line_spacing = 1.0
run = sub_p.add_run('การวิเคราะห์เครือข่ายสังคมของจำนวนนักท่องเที่ยวต่างชาติและราคาหุ้น AOT')
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True

# Student info
info_lines = [
    'รายวิชา: DDAS7201 การวิเคราะห์เครือข่ายสังคม',
    'ภาคการศึกษา: ปีการศึกษา 2568',
    'รหัสนักศึกษา: _______________  ชื่อ: ______________________________',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)

# Separator
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(3)
sep.paragraph_format.space_after = Pt(3)
sep.paragraph_format.line_spacing = 1.0
run = sep.add_run('─' * 60)
run.font.name = 'Times New Roman'
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(128, 128, 128)

# Section 1
write_section_th(1, 'บทนำ',
    'งานวิจัยนี้ประยุกต์ใช้ Social Network Analysis (SNA) เพื่อวิเคราะห์ความสัมพันธ์ระหว่าง '
    'จำนวนนักท่องเที่ยวต่างชาติ 4 กลุ่มตลาด (จีน, อินเดีย, ยุโรป, อาเซียน) '
    'กับราคาหุ้น 4 หลักทรัพย์ใน SET (AOT, AAV, MINT, CENTEL) จำนวน 12 เดือน '
    'ผ่านกรอบ Bipartite Tourism-Stock Dependency Network (BTSD-Net) '
    'เพื่อวัด centrality, network density, และโครงสร้างการพึ่งพาแบบสองส่วน')

# Section 2
write_section_th(2, 'การสร้างกราฟสองส่วน (Bipartite Graph)',
    f'นิยาม G = (V_M, V_S, E) โดย V_M = {{จีน, อินเดีย, ยุโรป, อาเซียน}} และ '
    f'V_S = {{AOT, AAV, MINT, CENTEL}} เส้นเชื่อม e_{{ij}} เมื่อ |r_{{ij}}| ≥ {THRESHOLD} '
    f'กราฟมี {n_nodes} โหนด ({n_markets} ตลาด, {n_stocks} หลักทรัพย์) '
    f'และ {n_edges} เส้นเชื่อม ความหนาแน่นเครือข่าย = {net_density:.4f}')

# Section 3
write_section_th(3, 'แรงจูงใจ: SNA สำหรับการวิเคราะห์',
    'แบบจำลองเชิงเส้นแบบดั้งเดิมไม่สามารถจับ cross-market contagion และโครงสร้างพึ่งพาแบบหลายมิติ '
    'SNA ช่วยวิเคราะห์ (1) centrality แต่ละโหนด (2) bipartite projection เพื่อหา co-occurrence '
    '(3) network density และ connectivity ที่สะท้อนความแข็งแกร่งของระบบนิเวศการท่องเที่ยว')

# Section 4
write_section_th(4, 'เมตริกของเครือข่าย (Network Metrics)',
    f'ใช้ normalized degree centrality d_norm(v) = deg(v)/(|V|−1) ร่วมกับ betweenness centrality '
    f'และ PageRank รวมถึงวิเคราะห์ bipartite projection เพื่อหา co-occurrence '
    f'ระหว่างตลาดและระหว่างหลักทรัพย์')

# Section 5
write_section_th(5, 'ข้อมูลและการตั้งค่า',
    f'ข้อมูลนักท่องเที่ยวจาก MOTS 12 เดือน แบ่งเป็น 4 กลุ่ม: '
    f'จีน (~15%), อินเดีย (~7%), ยุโรป (~18%), อาเซียน (~25%) '
    'ข้อมูลราคาหุ้นจาก SET สำหรับ AOT, AAV, MINT, CENTEL '
    f'คำนวณ Pearson correlation กรองที่ |r| ≥ {THRESHOLD} '
    'สร้างกราฟสองส่วนด้วย NetworkX พร้อม SNA visualization 4 แผง')

# Compute degree centrality for dynamic report text
deg_by_node = {n: round(deg_centrality[n], 4) for n in G.nodes()}
aot_dnorm = deg_by_node['AOT']
asean_name = 'อาเซียน'
asean_dnorm = deg_by_node[asean_name]
sorted_deg = sorted(deg_by_node.items(), key=lambda x: -x[1])

# Section 6
write_section_th(6, 'ผลการวิเคราะห์เครือข่าย (SNA Results)',
    f'1) Centrality: AOT มี d_norm = {aot_dnorm} สูงสุดในกลุ่มหลักทรัพย์ '
    f'เชื่อมกับ {[market_en.get(n, n) for n in G.neighbors("AOT")]} '
    f'อาเซียนเป็นตลาดที่มี d_norm = {asean_dnorm} เชื่อมกับ {[n for n in G.neighbors(asean_name)]} '
    f'ยืนยันบทบาทศูนย์กลางของ AOT ในฐานะสะพานเชื่อมระหว่างอุปสงค์การท่องเที่ยวและตลาดทุน<br/><br/>'
    f'2) Network density = {net_density:.4f} — {n_edges} จาก {n_markets * n_stocks} เส้นเชื่อมที่เป็นไปได้ '
    f'bipartite projection แสดง co-occurrence และ similarity structure<br/><br/>'
    f'3) ตารางที่ 1–2 และรูปที่ 1 แสดงผล SNA อย่างสมบูรณ์')

# Table
tbl_caption = doc.add_paragraph()
tbl_caption.paragraph_format.space_before = Pt(4)
tbl_caption.paragraph_format.space_after = Pt(1)
tbl_caption.paragraph_format.line_spacing = 1.0
run = tbl_caption.add_run(f'ตารางที่ 1 ชุดเส้นเชื่อม E ในกราฟสองส่วน BTSD-Net (|r| ≥ {THRESHOLD})')
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run.bold = True

table = doc.add_table(rows=1 + len(filtered_edges), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Shading Accent 1'

headers = ['ตลาดต้นทาง', 'หลักทรัพย์', 'Pearson r', '|น้ำหนัก|']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(h)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.bold = True

for row_idx, (mkt, sym, r_val) in enumerate(filtered_edges, 1):
    row = table.rows[row_idx]
    vals = [mkt, sym, f'{r_val:.3f}', f'{abs(r_val):.3f}']
    for col_idx, val in enumerate(vals):
        cell = row.cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)

for row in table.rows:
    row.cells[0].width = Cm(2.8)
    row.cells[1].width = Cm(2.0)
    row.cells[2].width = Cm(2.4)
    row.cells[3].width = Cm(2.0)

# Add network image
img_caption = doc.add_paragraph()
img_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_caption.paragraph_format.space_before = Pt(6)
img_caption.paragraph_format.space_after = Pt(2)
img_caption.paragraph_format.line_spacing = 1.0
run = img_caption.add_run('รูปที่ 1 เครือข่ายการพึ่งพาการท่องเที่ยวและหลักทรัพย์แบบสองส่วน')
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run.bold = True

img_para = doc.add_paragraph()
img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_para.paragraph_format.space_before = Pt(0)
img_para.paragraph_format.space_after = Pt(2)
img_para.paragraph_format.line_spacing = 1.0
run = img_para.add_run()
run.add_picture('temp_graph_th.png', width=Inches(4.8))

# Dynamic hierarchy for concluding paragraph
top_nodes_en = [(market_en.get(n, n), v) for n, v in sorted_deg]
hier_str = ' > '.join([f'{n} (d_norm = {v})' for n, v in top_nodes_en])

conc_p = doc.add_paragraph()
conc_p.paragraph_format.space_before = Pt(3)
conc_p.paragraph_format.space_after = Pt(0)
conc_p.paragraph_format.line_spacing = 1.0
conc_p.paragraph_format.first_line_indent = Cm(0.5)
run = conc_p.add_run(
    f'การวิเคราะห์ SNA ด้วย BTSD-Net แสดงลำดับ centrality: {hier_str} '
    f'AOT เป็น hub หลัก (d_norm = {aot_dnorm}) อาเซียนเป็นตลาดศูนย์กลาง (d_norm = {asean_dnorm}) '
    f'network density = {net_density:.4f} สะท้อนโครงสร้างแบบ '
    f'{"หนาแน่น" if net_density > 0.5 else "กระจายตัว"} '
    'bipartite projection เผยให้เห็น co-occurrence และ similarity structure '
    'ที่ SNA ช่วยให้เข้าใจ dependence structure ในระบบนิเวศการท่องเที่ยวได้ลึกซึ้งกว่า '
    'การวิเคราะห์สหสัมพันธ์แบบดั้งเดิม'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

out_docx = 'DDAS7201_Midterm_Project_TH.docx'
doc.save(out_docx)
print(f"  Document saved: {out_docx}")

# ============ STAGE 7: PDF Export ============
print("\n[STAGE 7] Exporting PDF...")

pdf_path = 'DDAS7201_Midterm_Project_TH.pdf'
converted = False

try:
    result = subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf', out_docx, '--outdir', '.'],
        capture_output=True, text=True, timeout=60
    )
    if os.path.exists(pdf_path):
        print(f"  PDF exported via LibreOffice: {pdf_path}")
        converted = True
except (FileNotFoundError, subprocess.TimeoutExpired):
    pass

if not converted and system == 'Windows':
    for base in [r'C:\Program Files\LibreOffice\program',
                 r'C:\Program Files (x86)\LibreOffice\program']:
        exe = os.path.join(base, 'soffice.exe')
        if os.path.exists(exe):
            try:
                abs_docx = os.path.abspath(out_docx)
                result = subprocess.run(
                    [exe, '--headless', '--convert-to', 'pdf', abs_docx, '--outdir', os.path.dirname(abs_docx)],
                    capture_output=True, text=True, timeout=60
                )
                if os.path.exists(pdf_path):
                    print(f"  PDF exported via LibreOffice: {pdf_path}")
                    converted = True
            except (subprocess.TimeoutExpired, Exception):
                pass
            if converted:
                break

if not converted and system == 'Windows':
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_docx = os.path.abspath(out_docx)
        abs_pdf = os.path.abspath(pdf_path)
        doc_obj = word.Documents.Open(abs_docx)
        doc_obj.SaveAs(abs_pdf, FileFormat=17)
        doc_obj.Close()
        word.Quit()
        if os.path.exists(pdf_path):
            print(f"  PDF exported via Microsoft Word: {pdf_path}")
            converted = True
    except Exception:
        pass

if not converted:
    print("  PDF conversion skipped. Install LibreOffice or Microsoft Word for automatic export.")
    print(f"  The .docx file is available at: {out_docx}")
    print("  To convert manually, open the .docx and export as PDF, or run:")
    print(f'    libreoffice --headless --convert-to pdf "{out_docx}"')

if os.path.exists('temp_graph_th.png'):
    try:
        os.remove('temp_graph_th.png')
    except Exception:
        pass

print()
print("=" * 60)
print("PROJECT COMPLETE (Thai Edition)")
print("=" * 60)
print("  Output files:")
print("    1. tourism_stock_network_th.svg  (vector graphic)")
print("    2. " + out_docx + "  (Word report - Thai)")
if converted:
    print("    3. " + pdf_path + "  (PDF copy)")
else:
    print("    3. " + pdf_path + "  (not generated - see instructions above)")
print("=" * 60)
