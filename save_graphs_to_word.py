#!/usr/bin/env python3
"""Standalone: generate SNA graphs + save to Word (no Streamlit needed)"""
import subprocess, sys, os, warnings, io, csv, urllib.request, datetime, platform
warnings.filterwarnings('ignore')
for pkg in ['python-docx','networkx','matplotlib','pandas','numpy','scipy','yfinance']:
    try:
        __import__(pkg.replace('-', '_'))
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg, '-q'])

import numpy as np
import pandas as pd
import networkx as nx
import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.dates as mdates
from scipy.stats import spearmanr
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import yfinance as yf

print("Loading data...")
# MOTS
url = ("https://ckan.mots.go.th/dataset/445c66d8-a06a-49d9-adfc-35faca6fc785/"
       "resource/faffc63c-9507-451a-80b7-554cc0787368/download/est_2024_04_01.csv")
req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
raw = urllib.request.urlopen(req, timeout=60).read()
content = raw.decode('utf-8-sig')
reader = csv.DictReader(io.StringIO(content))
rows = list(reader)
ASEAN = {'BRUNEI','CAMBODIA','INDONESIA','LAOS','MALAYSIA','MYANMAR','PHILIPPINES','SINGAPORE','THAILAND','VIETNAM','TIMOR-LESTE','VIET NAM'}
records = []
for row in rows:
    country = row['Country'].strip().upper()
    continent = row['continent'].strip()
    if 'CHINA' in country or 'HONG KONG' in country or 'TAIWAN' in country or 'MACAO' in country:
        mkt = 'China'
    elif 'INDIA' in country: mkt = 'India'
    elif continent == 'Europe': mkt = 'Europe'
    elif country in ASEAN: mkt = 'ASEAN'
    else: continue
    try: num = int(row[' Number '].replace(',', '').strip())
    except: continue
    parts = row['date'].strip().split('/')
    if len(parts) == 3:
        d, m, y = int(parts[0]), int(parts[1]), int(parts[2])
        records.append({'date': datetime.date(y, m, d), 'market': mkt, 'arrivals': num})
df = pd.DataFrame(records)
df['month_key'] = df['date'].apply(lambda d: datetime.date(d.year, d.month, 1))
monthly = df.groupby(['market', 'month_key'])['arrivals'].sum().reset_index()
pivot = monthly.pivot_table(index='month_key', columns='market', values='arrivals', aggfunc='sum').fillna(0).sort_index()
pivot.index = pd.to_datetime(pivot.index)

# Stocks
stock_prices = pd.DataFrame(index=pivot.index)
for sym, label in zip(['AOT.BK','AAV.BK','MINT.BK','CENTEL.BK'], ['AOT','AAV','MINT','CENTEL']):
    df = yf.download(sym, start=pivot.index[0].strftime('%Y-%m-%d'),
                     end=pivot.index[-1].strftime('%Y-%m-%d'), auto_adjust=True, progress=False)
    if df.empty: continue
    df.columns = [c[0] if isinstance(c, tuple) else c for c in df.columns]
    mc = df['Close'].resample('ME').last()
    mc.index = mc.index.tz_localize(None)
    aligned = mc.reindex(pivot.index, method='ffill')
    stock_prices[label] = aligned.values

# Build graph
THRESHOLD = 0.20
market_names = ['China','India','Europe','ASEAN']
stock_labels = ['AOT','AAV','MINT','CENTEL']
G = nx.Graph()
for m in market_names: G.add_node(m, bipartite=0, type='market')
for s in stock_labels: G.add_node(s, bipartite=1, type='stock')
corr_matrix = np.zeros((4, 4))
edges = []
for i, mkt in enumerate(market_names):
    for j, stk in enumerate(stock_labels):
        r, _ = spearmanr(pivot[mkt].values, stock_prices[stk].values)
        corr_matrix[i, j] = r
        if abs(r) >= THRESHOLD:
            G.add_edge(mkt, stk, weight=r, weight_abs=abs(r))
            edges.append((mkt, stk, r))
print(f"  Graph: {G.number_of_nodes()} nodes, {G.number_of_edges()} edges")

# Metrics
deg = nx.degree_centrality(G)
close = nx.closeness_centrality(G)
btwn = nx.betweenness_centrality(G) if nx.is_connected(G) else {}
if not btwn:
    for comp in nx.connected_components(G):
        btwn.update(nx.betweenness_centrality(G.subgraph(comp)))
pr = nx.pagerank(G, alpha=0.85, weight='weight_abs')
bottom, top = nx.bipartite.sets(G)
net_density = nx.density(G)

# Font
system = platform.system()
candidates = {'Windows': ['Tahoma','Cordia New','Angsana New','Leelawadee'],
              'Linux': ['Noto Sans Thai','Norasi','Laksaman','Waree'],
              'Darwin': ['Thonburi','Tahoma']}.get(system, ['DejaVu Sans','Tahoma'])
avail = {f.name for f in fm.fontManager.ttflist}
tf = next((c for c in candidates if c in avail), 'DejaVu Sans')
plt.rcParams['font.family'] = tf; plt.rcParams['axes.unicode_minus'] = False

m_colors = ['#E8630A','#F4A261','#D62828','#E9C46A']
s_colors = {'AOT':'#003049','AAV':'#1D70B8','MINT':'#2A9D8F','CENTEL':'#7B2CBF'}

print("Generating SNA figure...")
m_nodes = sorted(bottom, key=lambda n: market_names.index(n))
s_nodes = sorted(top, key=lambda n: stock_labels.index(n))

fig, axes = plt.subplots(2, 2, figsize=(15, 10))
ax_bp, ax_proj, ax_mm, ax_deg = axes[0,0], axes[0,1], axes[1,0], axes[1,1]

pos = nx.bipartite_layout(G, m_nodes, align='vertical', scale=1.0)
for n in pos:
    pos[n] = (-0.7, pos[n][1]) if n in m_nodes else (0.7, pos[n][1])
ns = [500 + deg[n] * 3000 for n in G.nodes()]
nc = [m_colors[m_nodes.index(n) % 4] if n in m_nodes else s_colors[n] for n in G.nodes()]
ew = [G[u][v]['weight_abs'] * 3.5 for u, v in G.edges()]
nx.draw_networkx_edges(G, pos, ax=ax_bp, width=ew, alpha=0.55, edge_color='#555555')
nx.draw_networkx_nodes(G, pos, ax=ax_bp, node_size=ns, node_color=nc, edgecolors='white', linewidths=2.0, alpha=0.92)
nx.draw_networkx_labels(G, pos, ax=ax_bp, font_size=10, font_color='white', font_weight='bold')
el = {(u, v): f"r={G[u][v]['weight']:.2f}" for u, v in G.edges()}
nx.draw_networkx_edge_labels(G, pos, edge_labels=el, ax=ax_bp, font_size=6.5, label_pos=0.5, bbox=dict(alpha=0.7, ec='none', pad=2))
ax_bp.set_title('(a) Bipartite Tourism-Stock Network', fontsize=12, fontweight='bold'); ax_bp.axis('off')

G_ss = nx.bipartite.weighted_projected_graph(G, top)
if G_ss.number_of_edges():
    pss = nx.circular_layout(G_ss)
    sw = [G_ss[u][v]['weight'] for u, v in G_ss.edges()]
    mw = max(sw) if sw else 1
    nx.draw(G_ss, pss, ax=ax_proj, node_size=[800+deg[n]*2000 for n in G_ss.nodes()],
            node_color=[s_colors[n] for n in G_ss.nodes()], edgecolors='white', linewidths=2.0,
            width=[w/mw*5 for w in sw], alpha=0.85, edge_color='#555555')
    nx.draw_networkx_labels(G_ss, pss, ax=ax_proj, font_size=10, font_color='white', font_weight='bold')
    nx.draw_networkx_edge_labels(G_ss, pss, {(u,v):f"{G_ss[u][v]['weight']:.1f}" for u,v in G_ss.edges()}, ax=ax_proj, font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
ax_proj.set_title('(b) Stock Co-occurrence Projection', fontsize=12, fontweight='bold'); ax_proj.axis('off')

G_mm = nx.bipartite.weighted_projected_graph(G, bottom)
if G_mm.number_of_edges():
    pmm = nx.circular_layout(G_mm)
    mw2 = [G_mm[u][v]['weight'] for u, v in G_mm.edges()]
    mwm = max(mw2) if mw2 else 1
    mm_nodes = list(G_mm.nodes())
    mm_colors = [m_colors[mm_nodes.index(n) % 4] for n in mm_nodes]
    nx.draw(G_mm, pmm, ax=ax_mm, node_size=[600+deg[n]*2000 for n in mm_nodes],
            node_color=mm_colors, edgecolors='white', linewidths=2.0,
            width=[w/mwm*5 for w in mw2], alpha=0.85, edge_color='#555555')
    nx.draw_networkx_labels(G_mm, pmm, ax=ax_mm, font_size=8)
    nx.draw_networkx_edge_labels(G_mm, pmm, {(u,v):f"{G_mm[u][v]['weight']:.1f}" for u,v in G_mm.edges()}, ax=ax_mm, font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
ax_mm.set_title('(c) Market Co-occurrence Projection', fontsize=12, fontweight='bold'); ax_mm.axis('off')

metrics = pd.DataFrame({'Node': list(deg.keys()), 'd_norm': [round(deg[n],4) for n in deg.keys()]})
sm = metrics.sort_values('d_norm', ascending=True)
bc = [m_colors[m_nodes.index(n)%4] if n in m_nodes else s_colors.get(n,'#999') for n in sm['Node']]
ax_deg.barh(range(len(sm)), sm['d_norm'].values, color=bc, edgecolor='white', height=0.6)
ax_deg.set_yticks(range(len(sm))); ax_deg.set_yticklabels(sm['Node'].values, fontsize=9)
ax_deg.set_xlabel('d_norm'); ax_deg.set_title('(d) Node Centrality Ranking', fontsize=12, fontweight='bold')
ax_deg.grid(True, alpha=0.3, axis='x')
for i, v in enumerate(sm['d_norm'].values): ax_deg.text(v+0.01, i, f'{v:.4f}', va='center', fontsize=8)

from matplotlib.patches import Patch
fig.legend(handles=[
    Patch(color='#E8630A',label='China'),Patch(color='#F4A261',label='India'),
    Patch(color='#D62828',label='Europe'),Patch(color='#E9C46A',label='ASEAN'),
    Patch(color='#003049',label='AOT'),Patch(color='#1D70B8',label='AAV'),
    Patch(color='#2A9D8F',label='MINT'),Patch(color='#7B2CBF',label='CENTEL')],
    loc='lower center', ncol=8, fontsize=8, bbox_to_anchor=(0.5, -0.02))
plt.tight_layout(); plt.subplots_adjust(bottom=0.08)
fig.savefig('_sna_fig.png', dpi=200, bbox_inches='tight')
plt.close()

print("Generating time series figure...")
fig2, ax2 = plt.subplots(figsize=(12, 5))
ts_colors = {'China':'#E8630A','India':'#F4A261','Europe':'#D62828','ASEAN':'#E9C46A'}
for m in market_names:
    ax2.plot(pivot.index, pivot[m].values/1e6, color=ts_colors[m], linewidth=1.5, label=m)
ax2.set_title('Monthly Tourist Arrivals (2015-2023)', fontsize=13, fontweight='bold')
ax2.set_ylabel('Arrivals (millions)'); ax2.set_xlabel('Year')
ax2.legend(fontsize=9); ax2.xaxis.set_major_locator(mdates.YearLocator(2))
ax2.xaxis.set_major_formatter(mdates.DateFormatter('%Y'))
plt.setp(ax2.xaxis.get_majorticklabels(), rotation=45, ha='right')
ax2.grid(True, alpha=0.3); plt.tight_layout()
fig2.savefig('_ts_fig.png', dpi=200, bbox_inches='tight')
plt.close()

print("Building Word document...")
doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(1.27); sec.bottom_margin = Cm(1.27)
    sec.left_margin = Cm(1.27); sec.right_margin = Cm(1.27)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(10)
style.paragraph_format.space_before = Pt(0); style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.0

def add_heading_styled(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.font.color.rgb = RGBColor(0,0,0)
    h.paragraph_format.space_before = Pt(5); h.paragraph_format.space_after = Pt(1); h.paragraph_format.line_spacing = 1.0

tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('BTSD-Net SNA Report'); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run(f'Tourism-Stock Dependency Network | threshold = {THRESHOLD}'); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
doc.add_paragraph()

add_heading_styled('1. SNA 4-Panel Visualization')
ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER; ip.add_run().add_picture('_sna_fig.png', width=Inches(5.8))

add_heading_styled('2. Tourist Arrivals Time Series')
ip2 = doc.add_paragraph(); ip2.alignment = WD_ALIGN_PARAGRAPH.CENTER; ip2.add_run().add_picture('_ts_fig.png', width=Inches(5.5))

add_heading_styled('3. Network Summary')
p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.0
r = p.add_run(f'Nodes: {G.number_of_nodes()} ({len(bottom)} markets, {len(top)} stocks)\nEdges: {G.number_of_edges()}\nDensity: {net_density:.4f}\nConnected: {"Yes" if nx.is_connected(G) else "No"}'); r.font.name = 'Times New Roman'; r.font.size = Pt(10)

add_heading_styled('4. Node Centrality Table')
tbl = doc.add_table(rows=1+len(deg), cols=5); tbl.style = 'Light Shading Accent 1'
for ci, h in enumerate(['Node','Type','d_norm','Betweenness','PageRank']):
    c = tbl.rows[0].cells[ci]; c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = 'Times New Roman'; r.font.size = Pt(8); r.bold = True
for ri, n in enumerate(sorted(deg, key=lambda x:-deg[x])):
    for ci, v in enumerate([n, 'Market' if n in bottom else 'Stock', f'{deg[n]:.4f}', f'{btwn.get(n,0):.4f}', f'{pr.get(n,0):.4f}']):
        c = tbl.rows[ri+1].cells[ci]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v); r.font.name = 'Times New Roman'; r.font.size = Pt(8)

add_heading_styled('5. Filtered Edges')
tbl2 = doc.add_table(rows=1+len(edges), cols=3); tbl2.style = 'Light Shading Accent 1'
for ci, h in enumerate(['Market','Stock','Spearman r']):
    c = tbl2.rows[0].cells[ci]; c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = 'Times New Roman'; r.font.size = Pt(8); r.bold = True
for ri, (mkt, stk, rv) in enumerate(edges, 1):
    for ci, v in enumerate([mkt, stk, f'{rv:.3f}']):
        c = tbl2.rows[ri].cells[ci]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v); r.font.name = 'Times New Roman'; r.font.size = Pt(8)

out = 'BTSDNet_SNA_Report.docx'
doc.save(out)
os.remove('_sna_fig.png'); os.remove('_ts_fig.png')
print(f"\nDone! Saved: {out}")
