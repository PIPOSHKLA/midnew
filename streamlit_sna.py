#!/usr/bin/env python3
"""Streamlit SNA Dashboard + Word export"""
import subprocess, sys, os, warnings, io, csv, urllib.request, datetime
warnings.filterwarnings('ignore')
for pkg in ['streamlit', 'python-docx', 'networkx', 'matplotlib', 'pandas', 'numpy', 'scipy', 'yfinance']:
    try:
        __import__(pkg.replace('-', '_'))
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg, '-q'])

import streamlit as st
import numpy as np
import pandas as pd
import networkx as nx
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.dates as mdates
from scipy.stats import spearmanr
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import yfinance as yf
import platform

st.set_page_config(page_title="BTSD-Net SNA Dashboard", layout="wide")
st.title("Bipartite Tourism-Stock Dependency Network")
st.markdown("SNA Dashboard — DDAS7201 Midterm Project")

# --- Helper functions ---
@st.cache_data
def load_data():
    ASEAN = {'BRUNEI','CAMBODIA','INDONESIA','LAOS','MALAYSIA','MYANMAR','PHILIPPINES','SINGAPORE','THAILAND','VIETNAM','TIMOR-LESTE','VIET NAM'}
    url = ("https://ckan.mots.go.th/dataset/445c66d8-a06a-49d9-adfc-35faca6fc785/"
           "resource/faffc63c-9507-451a-80b7-554cc0787368/download/est_2024_04_01.csv")
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    raw = urllib.request.urlopen(req, timeout=60).read()
    content = raw.decode('utf-8-sig')
    reader = csv.DictReader(io.StringIO(content))
    records = []
    for row in reader:
        country = row['Country'].strip().upper()
        continent = row['continent'].strip()
        if 'CHINA' in country or 'HONG KONG' in country or 'TAIWAN' in country or 'MACAO' in country:
            mkt = 'China'
        elif 'INDIA' in country:
            mkt = 'India'
        elif continent == 'Europe':
            mkt = 'Europe'
        elif country in ASEAN:
            mkt = 'ASEAN'
        else:
            continue
        try:
            num = int(row[' Number '].replace(',', '').strip())
        except:
            continue
        parts = row['date'].strip().split('/')
        if len(parts) == 3:
            d, m, y = int(parts[0]), int(parts[1]), int(parts[2])
            dt = datetime.date(y, m, d)
            records.append({'date': dt, 'market': mkt, 'arrivals': num})
    df = pd.DataFrame(records)
    df['month_key'] = df['date'].apply(lambda d: datetime.date(d.year, d.month, 1))
    monthly = df.groupby(['market', 'month_key'])['arrivals'].sum().reset_index()
    pivot = monthly.pivot_table(index='month_key', columns='market', values='arrivals', aggfunc='sum').fillna(0).sort_index()
    pivot.index = pd.to_datetime(pivot.index)
    for m in ['China','India','Europe','ASEAN']:
        if m not in pivot.columns: pivot[m] = 0.0
    return pivot

@st.cache_data
def load_stocks(pivot):
    stock_prices = pd.DataFrame(index=pivot.index)
    for sym, label in zip(['AOT.BK','AAV.BK','MINT.BK','CENTEL.BK'], ['AOT','AAV','MINT','CENTEL']):
        df = yf.download(sym, start=pivot.index[0].strftime('%Y-%m-%d'),
                         end=pivot.index[-1].strftime('%Y-%m-%d'), auto_adjust=True, progress=False)
        if df.empty: continue
        df.columns = [c[0] if isinstance(c, tuple) else c for c in df.columns]
        mc = df['Close'].resample('ME').last()
        mc.index = mc.index.tz_localize(None)
        aligned = mc.reindex(pivot.index, method='ffill')
        stock_prices[label] = aligned.values
    return stock_prices

def build_graph(pivot, stock_prices, threshold):
    market_names = ['China','India','Europe','ASEAN']
    stock_labels = ['AOT','AAV','MINT','CENTEL']
    corr_matrix = np.zeros((4, 4))
    for i, mkt in enumerate(market_names):
        for j, stk in enumerate(stock_labels):
            r, _ = spearmanr(pivot[mkt].values, stock_prices[stk].values)
            corr_matrix[i, j] = r
    G = nx.Graph()
    for m in market_names: G.add_node(m, bipartite=0, type='market')
    for s in stock_labels: G.add_node(s, bipartite=1, type='stock')
    edges = []
    for i, mkt in enumerate(market_names):
        for j, stk in enumerate(stock_labels):
            if abs(corr_matrix[i, j]) >= threshold:
                G.add_edge(mkt, stk, weight=corr_matrix[i, j], weight_abs=abs(corr_matrix[i, j]))
                edges.append((mkt, stk, corr_matrix[i, j]))
    return G, corr_matrix, edges

def compute_sna_metrics(G):
    deg = nx.degree_centrality(G)
    close = nx.closeness_centrality(G)
    if nx.is_connected(G):
        btwn = nx.betweenness_centrality(G)
    else:
        btwn = {}
        for comp in nx.connected_components(G):
            btwn.update(nx.betweenness_centrality(G.subgraph(comp)))
    pr = nx.pagerank(G, alpha=0.85, weight='weight_abs')
    bottom, top = nx.bipartite.sets(G)
    return deg, close, btwn, pr, bottom, top

def make_figure(G, pivot, stock_prices, deg, close, btwn, pr, bottom, top, threshold):
    system = platform.system()
    candidates = {'Windows': ['Tahoma','Cordia New','Angsana New','Leelawadee'],
                  'Linux': ['Noto Sans Thai','Norasi','Laksaman','Waree'],
                  'Darwin': ['Thonburi','Tahoma']}.get(system, ['DejaVu Sans','Tahoma'])
    avail = {f.name for f in fm.fontManager.ttflist}
    tf = next((c for c in candidates if c in avail), 'DejaVu Sans')
    plt.rcParams['font.family'] = tf
    plt.rcParams['axes.unicode_minus'] = False

    m_names = ['China','India','Europe','ASEAN']
    s_names = ['AOT','AAV','MINT','CENTEL']
    m_colors = ['#E8630A','#F4A261','#D62828','#E9C46A']
    s_colors = {'AOT':'#003049','AAV':'#1D70B8','MINT':'#2A9D8F','CENTEL':'#7B2CBF'}
    market_nodes_list = sorted(bottom, key=lambda n: m_names.index(n) if n in m_names else 99)
    stock_nodes_list = sorted(top, key=lambda n: s_names.index(n) if n in s_names else 99)

    fig, axes = plt.subplots(2, 2, figsize=(15, 10))
    ax_bp, ax_proj, ax_mm, ax_deg = axes[0,0], axes[0,1], axes[1,0], axes[1,1]

    # (a) Bipartite
    pos = nx.bipartite_layout(G, market_nodes_list, align='vertical', scale=1.0)
    for node in pos:
        x, y = pos[node]
        pos[node] = (-0.7, y) if node in market_nodes_list else (0.7, y)
    ns = [500 + deg[n] * 3000 for n in G.nodes()]
    nc = []
    for n in G.nodes():
        if n in market_nodes_list:
            nc.append(m_colors[market_nodes_list.index(n) % len(m_colors)])
        else:
            nc.append(s_colors[n])
    ew = [G[u][v]['weight_abs'] * 3.5 for u, v in G.edges()]
    nx.draw_networkx_edges(G, pos, ax=ax_bp, width=ew, alpha=0.55, edge_color='#555555')
    nx.draw_networkx_nodes(G, pos, ax=ax_bp, node_size=ns, node_color=nc, edgecolors='white', linewidths=2.0, alpha=0.92)
    nx.draw_networkx_labels(G, pos, ax=ax_bp, font_size=10, font_color='white', font_weight='bold')
    el = {(u, v): f"r={G[u][v]['weight']:.2f}" for u, v in G.edges()}
    nx.draw_networkx_edge_labels(G, pos, edge_labels=el, ax=ax_bp, font_size=6.5, label_pos=0.5, bbox=dict(alpha=0.7, ec='none', pad=2))
    ax_bp.set_title('(a) Bipartite Tourism-Stock Network', fontsize=12, fontweight='bold', pad=8)
    ax_bp.axis('off')

    # (b) Stock-Stock projection
    G_ss = nx.bipartite.weighted_projected_graph(G, top) if G.number_of_edges() else nx.Graph()
    if G_ss.number_of_edges():
        pss = nx.circular_layout(G_ss)
        sw = [G_ss[u][v]['weight'] for u, v in G_ss.edges()]
        mw = max(sw) if sw else 1
        sn = [800 + deg[n] * 2000 for n in G_ss.nodes()]
        sc = [s_colors[n] for n in G_ss.nodes()]
        nx.draw(G_ss, pss, ax=ax_proj, node_size=sn, node_color=sc, edgecolors='white', linewidths=2.0, width=[w/mw*5 for w in sw], alpha=0.85, edge_color='#555555')
        nx.draw_networkx_labels(G_ss, pss, ax=ax_proj, font_size=10, font_color='white', font_weight='bold')
        sl = {(u, v): f"{G_ss[u][v]['weight']:.1f}" for u, v in G_ss.edges()}
        nx.draw_networkx_edge_labels(G_ss, pss, edge_labels=sl, ax=ax_proj, font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
    ax_proj.set_title('(b) Stock Co-occurrence Projection', fontsize=12, fontweight='bold', pad=8)
    ax_proj.axis('off')

    # (c) Market-Market projection
    G_mm = nx.bipartite.weighted_projected_graph(G, bottom) if G.number_of_edges() else nx.Graph()
    if G_mm.number_of_edges():
        pmm = nx.circular_layout(G_mm)
        mw2 = [G_mm[u][v]['weight'] for u, v in G_mm.edges()]
        mwm = max(mw2) if mw2 else 1
        mn = [600 + deg[n] * 2000 for n in G_mm.nodes()]
        mc2 = [m_colors[list(G_mm.nodes()).index(n) % len(m_colors)] for n in G_mm.nodes()]
        nx.draw(G_mm, pmm, ax=ax_mm, node_size=mn, node_color=mc2, edgecolors='white', linewidths=2.0, width=[w/mwm*5 for w in mw2], alpha=0.85, edge_color='#555555')
        nx.draw_networkx_labels(G_mm, pmm, ax=ax_mm, font_size=8)
        ml = {(u, v): f"{G_mm[u][v]['weight']:.1f}" for u, v in G_mm.edges()}
        nx.draw_networkx_edge_labels(G_mm, pmm, edge_labels=ml, ax=ax_mm, font_size=7, bbox=dict(alpha=0.7, ec='none', pad=2))
    ax_mm.set_title('(c) Market Co-occurrence Projection', fontsize=12, fontweight='bold', pad=8)
    ax_mm.axis('off')

    # (d) Centrality bar
    metrics = pd.DataFrame({'Node': list(deg.keys()), 'd_norm': [round(deg[n], 4) for n in deg.keys()]})
    sm = metrics.sort_values('d_norm', ascending=True)
    bc2 = []
    for nd in sm['Node']:
        if nd in market_nodes_list:
            bc2.append(m_colors[market_nodes_list.index(nd) % len(m_colors)])
        else:
            bc2.append(s_colors.get(nd, '#999999'))
    ax_deg.barh(range(len(sm)), sm['d_norm'].values, color=bc2, edgecolor='white', height=0.6)
    ax_deg.set_yticks(range(len(sm)))
    ax_deg.set_yticklabels(sm['Node'].values, fontsize=9)
    ax_deg.set_xlabel('d_norm (Normalized Degree Centrality)', fontsize=10)
    ax_deg.set_title('(d) Node Centrality Ranking', fontsize=12, fontweight='bold')
    ax_deg.grid(True, alpha=0.3, axis='x')
    for i, v in enumerate(sm['d_norm'].values):
        ax_deg.text(v + 0.01, i, f'{v:.4f}', va='center', fontsize=8)

    from matplotlib.patches import Patch
    leg = [
        Patch(color='#E8630A', label='China'), Patch(color='#F4A261', label='India'),
        Patch(color='#D62828', label='Europe'), Patch(color='#E9C46A', label='ASEAN'),
        Patch(color='#003049', label='AOT'), Patch(color='#1D70B8', label='AAV'),
        Patch(color='#2A9D8F', label='MINT'), Patch(color='#7B2CBF', label='CENTEL'),
    ]
    fig.legend(handles=leg, loc='lower center', ncol=8, fontsize=8, framealpha=0.9, bbox_to_anchor=(0.5, -0.02))
    plt.tight_layout()
    plt.subplots_adjust(bottom=0.08)
    return fig

def make_time_series_figure(pivot):
    system = platform.system()
    candidates = {'Windows': ['Tahoma','Cordia New','Angsana New','Leelawadee'],
                  'Linux': ['Noto Sans Thai','Norasi','Laksaman','Waree'],
                  'Darwin': ['Thonburi','Tahoma']}.get(system, ['DejaVu Sans','Tahoma'])
    avail = {f.name for f in fm.fontManager.ttflist}
    tf = next((c for c in candidates if c in avail), 'DejaVu Sans')
    plt.rcParams['font.family'] = tf

    m_colors = {'China':'#E8630A','India':'#F4A261','Europe':'#D62828','ASEAN':'#E9C46A'}
    fig, ax = plt.subplots(figsize=(12, 5))
    for m in ['China','India','Europe','ASEAN']:
        ax.plot(pivot.index, pivot[m].values / 1e6, color=m_colors[m], linewidth=1.5, label=m)
    ax.set_title('Monthly Tourist Arrivals by Market (2015-2023)', fontsize=13, fontweight='bold')
    ax.set_ylabel('Arrivals (millions)')
    ax.set_xlabel('Year')
    ax.legend(fontsize=9, framealpha=0.9)
    ax.xaxis.set_major_locator(mdates.YearLocator(2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y'))
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha='right')
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    return fig

def save_to_word(fig, ts_fig, pivot, G, deg, btwn, pr, threshold, edges, out_path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.27); sec.bottom_margin = Cm(1.27)
        sec.left_margin = Cm(1.27); sec.right_margin = Cm(1.27)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0); style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run('BTSD-Net: Social Network Analysis Report')
    r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run(f'Tourism-Stock Dependency Network (threshold = {threshold})')
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True

    doc.add_paragraph()  # spacer

    # Save plots as temp files
    fig.savefig('_temp_sna.png', dpi=200, bbox_inches='tight', pad_inches=0.3)
    ts_fig.savefig('_temp_ts.png', dpi=200, bbox_inches='tight', pad_inches=0.3)

    # Section 1: Network visualization
    h = doc.add_heading('1. SNA Visualization', level=2)
    for run in h.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(11)
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture('_temp_sna.png', width=Inches(5.8))

    # Section 2: Time series
    h = doc.add_heading('2. Tourist Arrivals Time Series', level=2)
    for run in h.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(11)
    ip2 = doc.add_paragraph()
    ip2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip2.add_run().add_picture('_temp_ts.png', width=Inches(5.5))

    # Section 3: Metrics table
    h = doc.add_heading('3. Network Metrics', level=2)
    for run in h.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(11)
    net_density = nx.density(G)
    n_nodes = G.number_of_nodes()
    n_edges = G.number_of_edges()
    bottom, top = nx.bipartite.sets(G)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(
        f'Nodes: {n_nodes} ({len(bottom)} markets, {len(top)} stocks)\n'
        f'Edges: {n_edges}\n'
        f'Network density: {net_density:.4f}\n'
        f'Connected: {nx.is_connected(G)}\n'
        f'Threshold: |r| >= {threshold}'
    )
    r.font.name = 'Times New Roman'; r.font.size = Pt(10)

    # Centrality table
    doc.add_paragraph()
    tbl = doc.add_table(rows=1 + len(deg), cols=5)
    tbl.style = 'Light Shading Accent 1'
    headers = ['Node', 'Type', 'd_norm', 'Betweenness', 'PageRank']
    for ci, hdr in enumerate(headers):
        cell = tbl.rows[0].cells[ci]; cell.text = ''
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr); r.font.name = 'Times New Roman'; r.font.size = Pt(8); r.bold = True
    for ri, n in enumerate(sorted(deg.keys(), key=lambda x: -deg[x])):
        vals = [n, 'Market' if n in bottom else 'Stock',
                f'{deg[n]:.4f}', f'{btwn.get(n, 0):.4f}', f'{pr.get(n, 0):.4f}']
        for ci, v in enumerate(vals):
            cell = tbl.rows[ri + 1].cells[ci]; cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v); r.font.name = 'Times New Roman'; r.font.size = Pt(8)

    # Section 4: Edge list
    h = doc.add_heading('4. Filtered Edges', level=2)
    for run in h.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(11)
    tbl2 = doc.add_table(rows=1 + len(edges), cols=3)
    tbl2.style = 'Light Shading Accent 1'
    for ci, hdr in enumerate(['Market', 'Stock', 'Spearman r']):
        cell = tbl2.rows[0].cells[ci]; cell.text = ''
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr); r.font.name = 'Times New Roman'; r.font.size = Pt(8); r.bold = True
    for ri, (mkt, stk, rv) in enumerate(edges, 1):
        for ci, v in enumerate([mkt, stk, f'{rv:.3f}']):
            cell = tbl2.rows[ri].cells[ci]; cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v); r.font.name = 'Times New Roman'; r.font.size = Pt(8)

    doc.save(out_path)
    os.remove('_temp_sna.png')
    os.remove('_temp_ts.png')

# --- Main app ---
with st.spinner("Loading MOTS data..."):
    pivot = load_data()
with st.spinner("Loading stock data from Yahoo Finance..."):
    stock_prices = load_stocks(pivot)

st.sidebar.header("Parameters")
threshold = st.sidebar.slider("Correlation threshold (|r| ≥)", 0.05, 0.50, 0.20, 0.05)

with st.spinner("Building graph..."):
    G, corr_matrix, edges = build_graph(pivot, stock_prices, threshold)
    deg, close, btwn, pr, bottom, top = compute_sna_metrics(G)

tab1, tab2, tab3, tab4 = st.tabs(["SNA Visualization", "Time Series", "Metrics Table", "Correlation Matrix"])

with tab1:
    st.subheader("SNA 4-Panel Visualization")
    fig = make_figure(G, pivot, stock_prices, deg, close, btwn, pr, bottom, top, threshold)
    st.pyplot(fig)
    plt.close(fig)

with tab2:
    st.subheader("Tourist Arrivals Time Series")
    ts_fig = make_time_series_figure(pivot)
    st.pyplot(ts_fig)
    plt.close(ts_fig)

with tab3:
    st.subheader("Network Metrics")
    col1, col2, col3 = st.columns(3)
    col1.metric("Nodes", G.number_of_nodes())
    col2.metric("Edges", G.number_of_edges())
    col3.metric("Density", f"{nx.density(G):.4f}")
    st.dataframe(pd.DataFrame({
        'Node': list(deg.keys()),
        'Type': ['Market' if n in bottom else 'Stock' for n in deg.keys()],
        'd_norm': [round(deg[n], 4) for n in deg.keys()],
        'Closeness': [round(close[n], 4) for n in deg.keys()],
        'Betweenness': [round(btwn.get(n, 0), 4) for n in deg.keys()],
        'PageRank': [round(pr.get(n, 4), 4) for n in deg.keys()],
    }).sort_values('d_norm', ascending=False), use_container_width=True)

    st.subheader("Bipartite Projection")
    G_ss = nx.bipartite.weighted_projected_graph(G, top)
    st.write(f"**Stock-Stock:** {G_ss.number_of_edges()} edges, density = {nx.density(G_ss):.4f}")
    G_mm = nx.bipartite.weighted_projected_graph(G, bottom)
    st.write(f"**Market-Market:** {G_mm.number_of_edges()} edges, density = {nx.density(G_mm):.4f}")

    st.subheader("Filtered Edges")
    st.dataframe(pd.DataFrame(edges, columns=['Market', 'Stock', 'Spearman r']), use_container_width=True)

with tab4:
    st.subheader("Spearman Correlation Matrix")
    market_names = ['China','India','Europe','ASEAN']
    stock_labels = ['AOT','AAV','MINT','CENTEL']
    corr_display = pd.DataFrame(corr_matrix, index=market_names, columns=stock_labels)
    st.dataframe(corr_display.style.background_gradient(cmap='RdBu_r', vmin=-1, vmax=1), use_container_width=True)

# Export section
st.sidebar.markdown("---")
st.sidebar.subheader("Export to Word")
doc_name = st.sidebar.text_input("Filename", "BTSDNet_SNA_Report.docx")
if st.sidebar.button("Generate Word Report"):
    with st.spinner("Generating Word document..."):
        fig_exp = make_figure(G, pivot, stock_prices, deg, close, btwn, pr, bottom, top, threshold)
        ts_exp = make_time_series_figure(pivot)
        save_to_word(fig_exp, ts_exp, pivot, G, deg, btwn, pr, threshold, edges, doc_name)
        plt.close(fig_exp); plt.close(ts_exp)
    st.sidebar.success(f"Saved: {doc_name}")
    with open(doc_name, "rb") as f:
        st.sidebar.download_button("Download Word File", f, file_name=doc_name)

st.sidebar.markdown("---")
st.sidebar.info("**Data sources:** MOTS CKAN + Yahoo Finance\n\n**Period:** 2015–2023 (102 months)")
